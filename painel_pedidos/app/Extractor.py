# Arquivo: app/Extractor.py (VERSÃO FINALÍSSIMA, COM CORREÇÃO DE REGEX)

"""Módulo central para extração de dados e geração de documentos.

Este arquivo foi atualizado para mesclar as funcionalidades existentes do grupo
com as novas capacidades de análise de PDF com IA e geração direta de PDF.
Ele serve como a principal biblioteca de lógica de negócio para contratos.
"""

import re
import fitz
import spacy
import json
from io import BytesIO
from docx import Document
import datetime
import openpyxl
from typing import Dict, Any, Optional
from flask import render_template
from weasyprint import HTML

# --- Configuração do Modelo de IA (spaCy) ---
try:
    nlp = spacy.load("pt_core_news_md")
    print("[INFO] Modelo de NLP (pt_core_news_lg) carregado com sucesso.")
except OSError:
    print("[AVISO] Modelo 'pt_core_news_lg' não foi encontrado. Execute: python -m spacy download pt_core_news_lg")
    nlp = None

# ==============================================================================
# SEÇÃO 1: EXTRAÇÃO DE DADOS DE PDF
# ==============================================================================

def extrair_dados_do_contrato_por_tipo(pdf_bytes: bytes, tipo_analise: str = 'padrao') -> Optional[Dict[str, Any]]:
    """Função principal que orquestra a análise de um contrato em PDF."""
    texto = _extrair_texto_de_pdf_bytes(pdf_bytes)
    if not texto:
        return None
    if tipo_analise == 'sistema':
        return _extrair_com_regex(texto)
    else:
        return _extrair_com_nlp(texto)

def _extrair_texto_de_pdf_bytes(pdf_bytes: bytes) -> Optional[str]:
    """Lê os bytes de um arquivo PDF usando PyMuPDF (fitz)."""
    try:
        with fitz.open(stream=pdf_bytes, filetype="pdf") as doc:
            return "".join(pagina.get_text("text") for pagina in doc)
    except Exception as e:
        print(f"[ERRO] Falha ao extrair texto do PDF a partir dos bytes: {e}")
        return None

def _extrair_com_regex(texto: str) -> Dict[str, Any]:
    """Extrai dados usando Regex, otimizada para contratos do sistema.
    
    Esta versão contém as correções finais para extração de email e produtos.
    """
    flags = re.DOTALL | re.IGNORECASE
    dados = {
        "Contratante": {"Nome": "N/A", "CPF": "N/A", "Telefone": "N/A", "Email": "N/A", "RG": "N/A", "Endereco": "N/A"},
        "Data_do_Evento": "N/A", "Local_do_Evento": "N/A", "produtosContratadosJson": "[]",
        "Data_de_Pagamento": "N/A", "Valor_Total_do_Pedido": "N/A", "FormaDePagamento": "N/A",
        "Responsavel": "N/A", "Como nos conheceu": "N/A"
    }

    # Bloco Contratante (sem alterações, já estava robusto)
    secao_contratante = re.search(r"CONTRATANTE:([\s\S]*?)CONTRATADO:", texto, flags)
    if secao_contratante:
        texto_contratante = secao_contratante.group(1)
        dados["Contratante"]["Nome"] = (m.group(1).strip() if (m := re.search(r"Sr\(a\)\s*(.*?),\s*brasileiro", texto_contratante, flags)) else "N/A")
        dados["Contratante"]["RG"] = (m.group(1).strip() if (m := re.search(r"RG:\s*([\d.\s-]+?)", texto_contratante, flags)) else "N/A")
        dados["Contratante"]["CPF"] = (m.group(1).strip() if (m := re.search(r"CPF:\s*([\d.\s-]+?),", texto_contratante, flags)) else "N/A")
        dados["Contratante"]["Endereco"] = (m.group(1).strip() if (m := re.search(r"domiciliado\(a\) na (.*?) - Tel\.", texto_contratante, flags)) else "N/A")
        dados["Contratante"]["Telefone"] = (m.group(1).strip() if (m := re.search(r"Tel\.\s*(.*?)\.", texto_contratante, flags)) else "N/A")

    # === CORREÇÃO 1: REGEX DE EMAIL MAIS FLEXÍVEL ===
    # Esta nova regra busca por "E-" seguido opcionalmente por quebra de linha "mail:",
    # ou simplesmente "Email:", em qualquer lugar do texto.
    email_match = re.search(r"(?:E-\s*\n?mail|Email):\s*([\w.%+-]+@[\w.-]+\.[a-zA-Z]{2,})", texto, flags)
    if email_match:
        dados["Contratante"]["Email"] = email_match.group(1).strip()

    # === CORREÇÃO 2: REGEX DE PRODUTOS MAIS ROBUSTA ===
    # Isola o bloco da tabela de forma mais precisa.
    produtos_bloco = re.search(r'CLÁUSULA 1 - PRODUTOS CONTRATADOS([\s\S]*?)TOTAL:\s*R\$', texto, flags)
    if produtos_bloco:
        produtos_lista = []
        # A nova regra para encontrar os itens agora espera o "R$" antes dos valores,
        # tornando a captura do nome do produto muito mais precisa.
        itens = re.findall(r'^\s*(\d+)\s+(.*?)\s+(R\$\s*[\d,.]+)\s+(R\$\s*[\d,.]+)\s*$', produtos_bloco.group(1), re.MULTILINE)
        for item in itens:
            produtos_lista.append({
                'Quantidade': item[0].strip(),
                'Produto': item[1].strip(),
                # Remove o "R$" para salvar apenas o número, se desejar.
                'Valor Unitário': item[2].replace('R$', '').strip(),
                'Valor Total Item': item[3].replace('R$', '').strip()
            })
        if produtos_lista:
            dados["produtosContratadosJson"] = json.dumps(produtos_lista, ensure_ascii=False)
            
    # O restante das regras já estava funcionando bem.
    dados["Valor_Total_do_Pedido"] = (m.group(1).strip() if (m := re.search(r"TOTAL:\s*(R\$\s*[\d,.]+)", texto, flags)) else "N/A")
    dados["Data_de_Pagamento"] = (m.group(1).strip() if (m := re.search(r"pagos no dia\s*([\d/]+)", texto, flags)) else "N/A")
    dados["FormaDePagamento"] = (m.group(1).strip() if (m := re.search(r"pagos no dia\s*[\d/]+\s*(.*?)\.", texto, flags)) else "N/A")
    clausula_11 = re.search(r"O evento acontecerá no dia:\s*(.*?)\s*-\s*Local do evento:\s*(.*?)Como nos conheceu:", texto, flags)
    if clausula_11:
        dados["Data_do_Evento"] = clausula_11.group(1).strip()
        dados["Local_do_Evento"] = clausula_11.group(2).strip()
    dados["Responsavel"] = (m.group(1).strip() if (m := re.search(r"RESPONSÁVEL PELO CONTRATO:\s*(.*?)\s*\n", texto, flags)) else "N/A")
    
    return dados

def _extrair_com_nlp(texto: str) -> Dict[str, Any]:
    """Extrai dados usando IA (NLP) para contratos com layout desconhecido."""
    if not nlp: raise Exception("Modelo de linguagem spaCy não foi carregado.")
    doc = nlp(texto)
    dados = {"Contratante": { "Nome": "Não encontrado", "CPF": "N/A", "Telefone": "N/A", "Email": "N/A" }, "Data_do_Evento": "Não encontrado", "Local_do_Evento": "Não encontrado", "produtosContratadosJson": '[]', "Data_de_Pagamento": "Verificar no Doc.", "Valor_Total_do_Pedido": "Não encontrado", "FormaDePagamento": "Verificar no Doc."}
    pessoas = [ent.text for ent in doc.ents if ent.label_ == "PER"]
    locais = [ent.text for ent in doc.ents if ent.label_ == "LOC"]
    if pessoas: dados["Contratante"]["Nome"] = pessoas[0]
    if locais: dados["Local_do_Evento"] = locais[0]
    dados["Contratante"]["CPF"] = (m.group(1) if (m := re.search(r"(\d{3}\.\d{3}\.\d{3}-\d{2})", texto)) else "N/A")
    dados["Contratante"]["Telefone"] = (m.group(1) if (m := re.search(r"(\(?\d{2}\)?\s*\d{4,5}-?\d{4})", texto)) else "N/A")
    dados["Contratante"]["Email"] = (m.group(1) if (m := re.search(r"([\w.\-]+@[\w.\-]+)", texto)) else "N/A")
    dados["Valor_Total_do_Pedido"] = (m.group(1) if (m := re.search(r"(?:valor\s*total|preço\s*final)[\s\S]*?(R\$\s*[\d.,]+)", texto, re.IGNORECASE)) else "Não encontrado")
    dados["Data_do_Evento"] = (m.group(1) if (m := re.search(r"data\s*do\s*evento[:\s]*(\d{2}/\d{2}/\d{4})", texto, re.IGNORECASE)) else "Não encontrado")
    return dados

# ==============================================================================
# SEÇÃO 2: GERAÇÃO DE DOCUMENTOS (Funcionalidades Mescladas e Aprimoradas)
# ==============================================================================

def gerar_contrato_docx(dados: Dict[str, Any]) -> Optional[BytesIO]:
    """Gera um documento .docx de contrato."""
    try:
        document = Document()
        document.add_heading('Divinos Doces Finos', 0)
        document.add_paragraph('CONTRATO', style='Normal')
        document.add_paragraph()
        contratante = dados.get('Contratante', {})
        contratante_texto = document.add_paragraph()
        contratante_texto.add_run('CONTRATANTE: ').bold = True
        contratante_texto.add_run(f"Sr(a) {contratante.get('Nome', 'N/A')}, brasileiro(a), portador(a) da cédula de RG: {contratante.get('RG', 'N/A')} e CPF: {contratante.get('CPF', 'N/A')}, residente e domiciliado(a) na {contratante.get('Endereco', 'N/A')} - Tel. {contratante.get('Telefone', 'N/A')}.")
        contratado_texto = document.add_paragraph()
        contratado_texto.add_run('CONTRATADO: ').bold = True
        contratado_texto.add_run(f"Divinos Doces Finos, inscrito sob o CNPJ: 18.826.801/0001-76, com sede na Rua Curupacê, 392 Mooca, São Paulo SP representado pela sócia proprietária Damaris Talita Macedo, portador do RG: 30.315.655-7.")
        document.add_paragraph()
        document.add_heading('CLÁUSULA 1 - PRODUTOS CONTRATADOS', level=1)
        produtos = dados.get('Produtos Contratados', [])
        if produtos:
            tabela = document.add_table(rows=1, cols=4)
            tabela.style = 'Table Grid'
            hdr_cells = tabela.rows[0].cells
            hdr_cells[0].text, hdr_cells[1].text, hdr_cells[2].text, hdr_cells[3].text = 'Quantidade', 'Produto', 'Valor Unitário', 'Valor Total'
            for item in produtos:
                row_cells = tabela.add_row().cells
                row_cells[0].text, row_cells[1].text, row_cells[2].text, row_cells[3].text = str(item.get('Quantidade', '')), str(item.get('Produto', '')), str(item.get('Valor Unitário', '')), str(item.get('Valor Total Item', ''))
            document.add_paragraph(f"TOTAL: R$ {dados.get('Valor Total do Pedido', 'N/A')}")
        else:
            document.add_paragraph("Nenhum produto adicionado.")
        document.add_heading('CLÁUSULA 2 - VALOR E FORMA DE PAGAMENTO', level=1)
        document.add_paragraph(f"O valor total de R$ {dados.get('Valor Total do Pedido', 'N/A')} referente aos produtos acima citados, foram pagos no dia {dados.get('Data de Pagamento', 'N/A')} {dados.get('Forma de Pagamento', 'N/A')}.")
        document.add_heading('CLÁUSULA 3 - EMBALAGEM DOS DOCES - FORMINHAS', level=1)
        document.add_paragraph('Os doces finos são entregues em forminhas no formato caixeta, na cor branca, todos decorados e prontos para o consumo. Os brigadeiros serão entregues em forminhas na cor branca nº 5.')
        document.add_paragraph('Caso o CONTRATANTE opte por embalagens decorativas, o mesmo deverá enviar ao CONTRATADO com no máximo 15 dias de antecedência ao evento, que entregará os doces finos dentro das embalagens decoradas, prontos para o consumo. Após esse prazo não recebemos.')
        document.add_paragraph('Por haver um manejo especial nas forminhas no modelo de flor e um custo maior de compra de caixas para armazenamento dos doces, é cobrado uma taxa adicional de R$0,10 por unidade, como consta abaixo:')
        document.add_paragraph('ATÉ 100 DOCES + R$10,00 / ATÉ 200 DOCES + R$20,00 ATÉ 300 DOCES + R$30,00 / ATÉ 400 DOCES + R$40,00 ACIMA DE 500 DOCES + R$50,00 e assim sucessivamente')
        document.add_paragraph()
        document.add_heading('CLÁUSULA 4 - EMBALAGENS DOS BEM-CASADOS', level=1)
        document.add_paragraph('Os bem-casados são entregues em papel crepom crepe plus, com celofane e fita de cetim de 7mm, nas cores enviadas na tabela completa. Os papéis perolados da linha especial serão cobrados R$ 0,40 a mais por unidade e os papéis dourado, prata, tiffany e marsala serão cobrados R$ 0,20 a mais por unidade, por se tratar de um papel especial e com maior custo. Tudo está discriminado na tabela de cores.')
        document.add_paragraph('Caso o CONTRATANTE opte por incluir, medalhinhas, tercinhos, renda, juta, tag ou outro item decorativo, deverá consultar antecipadamente a disponibilidade e todos os itens são colados com cola quente. A entrega dos itens deverá ocorrer com no máximo 15 dias antes do evento. Após esse prazo não recebemos. Por haver um manejo especial dos itens, será cobrado uma taxa adicional de R$0,10, como consta abaixo: ATÉ 100 BEM-CASADOS + R$10,00 / ATÉ 200 BEM-CASADOS + R$20,00 ATÉ 300 BEM-CASADOS + R$30,00 / ATÉ 400 BEM-CASADOS + R$40,00 ACIMA DE 500 BEM-CASADOS + R$50,00 e assim sucessivamente')
        document.add_paragraph('Caso opte pela aplicação de dois ou mais itens, será cobrado o valor de cada item.')
        document.add_paragraph()
        document.add_heading('CLÁUSULA 5 - ALTERAÇÕES', level=1)
        document.add_paragraph('Não recebemos forminhas, modificações, alterações em contrato em hipótese alguma na semana do evento.')
        document.add_paragraph()
        document.add_heading('CLÁUSULA 6 - ADIÇÃO DE NOVOS ITENS', level=1)
        document.add_paragraph('Caso haja a necessidade do CONTRATANTE adicionar novos itens ao pedido fechado, o valor dos produtos será de acordo com o valor vigente no momento da adição, mesmo que o contrato tenha sido fechado com valores promocionais.')
        document.add_paragraph('A adição de produtos ocorre de acordo com a disponibilidade de agenda. Não havendo disponibilidade para novos produtos ou pedidos, não será possível a complementação.')
        document.add_paragraph()
        document.add_heading('CLÁUSULA 7 - RETIRADA OU SERVIÇO DE ENTREGA', level=1)
        document.add_paragraph(f"A entrega ou retirada dos itens acima, deverá ser definida pela CONTRATANTE até 15 dias antes do evento. Em caso de entrega será cobrada taxa de deslocamento de R$ 6,00 por km ou a taxa mínima de R$50,00 (sujeito a disponibilidade na data e horário desejados). Não fazemos entregas aos domingos e feriados. A retirada dos produtos ocorre de segunda-feira à sábado, das 9h às 16h30, mediante agendamento com o setor responsável, não havendo expediente aos domingos e feriados.")
        document.add_paragraph()
        document.add_heading('CLÁUSULA 8 - ARMAZENAMENTO', level=1)
        document.add_paragraph('Todos os doces e/ou bem-casados deverão, obrigatoriamente, ser armazenados em geladeira até o momento da montagem da mesa para o evento. Validade 3 a 5 dias em geladeira.')
        document.add_paragraph()
        document.add_heading('CLÁUSULA 9 - LOCAÇÃO (SE HOUVER)', level=1)
        document.add_paragraph('Caso haja locação de bolo cenográfico, o CONTRATANTE deverá deixar uma caução no valor de R$300,00 ou o valor em dinheiro, como forma de garantia. O bolo cenográfico sendo locado e deverá retornar nas mesmas condições, em até 4 dias após a data da retirada. Na devolução do bolo cenográfico, será devolvido o valor total. Em caso de avarias será cobrado R$ 100,00 por andar (dependendo do modelo) para refazer cada andar danificado. O CONTRATANTE deverá tomar todos os cuidados necessários como: não expor ao calor excessivo, água ou qualquer outro líquido, não deverá apertar, amassar, não deixar convidados colocarem as mãos e deverá ser transportado com cuidado, pegando somente pela base de madeira.')
        document.add_paragraph()
        document.add_heading('CLÁUSULA 10 - REMARCAÇÃO', level=1)
        document.add_paragraph('Em caso de REMARCAÇÃO de data do evento superior a 6 meses, será cobrado um reequilíbrio econômico e financeiro de 10% sobre o valor do contrato, a cada 6 meses de diferença da data marcada inicialmente.')
        document.add_paragraph()
        document.add_heading('CLÁUSULA 11 - DATA E LOCAL DO EVENTO', level=1)
        document.add_paragraph(f"O evento acontecerá no dia: {dados.get('Data do Evento', 'N/A')} - Local do evento: {dados.get('Local do Evento', 'N/A')}")
        document.add_paragraph(f"Como nos conheceu: {dados.get('Como nos conheceu', 'N/A')}")
        document.add_paragraph()
        document.add_heading('CLÁUSULA 12 - CANCELAMENTO', level=1)
        document.add_paragraph('A CONTRATANTE pagará multa de 30% do valor do contrato em caso de cancelamento. O CONTRATADO pagará multa de 100% do valor do contrato em caso de cancelamento.')
        document.add_paragraph()
        document.add_paragraph(f"RESPONSÁVEL PELO CONTRATO: {dados.get('Responsavel', 'N/A')}")
        document.add_paragraph(f"São Paulo, {dados.get('Data de Pagamento', 'N/A')}")
        document.add_paragraph()
        document.add_paragraph('CONTRATANTE', style='Normal').bold = True
        document.add_paragraph('______________________________', style='Normal')
        document.add_paragraph('CONTRATADO', style='Normal').bold = True
        document.add_paragraph('______________________________', style='Normal')
        
        doc_stream = BytesIO()
        document.save(doc_stream)
        doc_stream.seek(0)
        return doc_stream
    except Exception as e:
        print(f"[ERRO] Falha ao gerar contrato DOCX: {e}")
        return None

def gerar_contrato_pdf_direto(dados: Dict[str, Any]) -> Optional[BytesIO]:
    """Gera um PDF diretamente a partir de um template HTML usando WeasyPrint."""
    try:
        html_string = render_template("contrato_template.html", dados=dados)
        pdf_bytes = HTML(string=html_string).write_pdf()
        pdf_stream = BytesIO(pdf_bytes)
        pdf_stream.seek(0)
        return pdf_stream
    except Exception as e:
        print(f"[ERRO] Falha ao gerar contrato PDF com WeasyPrint: {e}")
        return None

def gerar_relatorio_entrega(dados: Dict[str, Any]) -> Optional[BytesIO]:
    """Gera um relatório de entrega em .docx."""
    try:
        document = Document()
        document.add_heading('RELATÓRIO DE ENTREGA', 0)
        contratante = dados.get('Contratante', {})
        data_evento = dados.get('Data do Evento', 'Não informada')
        local_evento = dados.get('Local do Evento', 'Não informado')
        document.add_paragraph(f"Nome do Cliente: {contratante.get('Nome', 'Não encontrado')}")
        document.add_paragraph(f"Data do Evento: {data_evento}")
        document.add_paragraph(f"Local do Evento: {local_evento}")
        document.add_paragraph(f"Data de Emissão: {datetime.datetime.now().strftime('%d/%m/%Y')}")
        document.add_paragraph("\nProdutos Contratados:")
        produtos = dados.get('Produtos Contratados', [])
        if produtos:
            tabela = document.add_table(rows=1, cols=4)
            tabela.style = 'Table Grid'
            hdr_cells = tabela.rows[0].cells
            hdr_cells[0].text, hdr_cells[1].text, hdr_cells[2].text, hdr_cells[3].text = 'Quantidade', 'Produto', 'Valor Unitário', 'Valor Total'
            for item in produtos:
                row_cells = tabela.add_row().cells
                row_cells[0].text, row_cells[1].text, row_cells[2].text, row_cells[3].text = str(item.get('Quantidade', '')), str(item.get('Produto', '')), str(item.get('Valor Unitário', '')), str(item.get('Valor Total Item', ''))
        else:
            document.add_paragraph("Nenhum produto encontrado.")
        document.add_paragraph(f"\nValor Total do Pedido: R$ {dados.get('Valor Total do Pedido', 'Não encontrado')}")
        document.add_paragraph("\n\n\nAssinaturas:\n")
        document.add_paragraph("______________________________\nResponsável pela Entrega")
        document.add_paragraph("\n\n")
        document.add_paragraph("______________________________\nResponsável pela Retirada")
        doc_stream = BytesIO()
        document.save(doc_stream)
        doc_stream.seek(0)
        return doc_stream
    except Exception as e:
        print(f"[ERRO] Falha ao salvar relatório de entrega: {e}")
        return None

def exportar_para_excel(dados: Dict[str, Any]) -> Optional[BytesIO]:
    """Exporta os dados para um arquivo Excel."""
    try:
        workbook = openpyxl.Workbook()
        sheet = workbook.active
        sheet.title = "Dados do Contrato"
        sheet['A1'] = "Campo"
        sheet['B1'] = "Informação Extraída"
        linha_atual = 2
        for chave, valor in dados.items():
            if chave == 'Produtos Contratados':
                continue
            if isinstance(valor, dict):
                for sub_chave, sub_valor in valor.items():
                    sheet[f'A{linha_atual}'] = f"{chave} - {sub_chave}"
                    sheet[f'B{linha_atual}'] = sub_valor
                    linha_atual += 1
            else:
                sheet[f'A{linha_atual}'] = chave
                sheet[f'B{linha_atual}'] = valor
                linha_atual += 1
        linha_atual += 2
        # Verifica se 'Produtos Contratados' está no dicionário e não está vazio
        produtos_contratados = dados.get('Produtos Contratados')
        if produtos_contratados:
            headers_produtos = list(produtos_contratados[0].keys())
            for col_idx, header in enumerate(headers_produtos, 1):
                sheet.cell(row=linha_atual, column=col_idx, value=header)
            linha_atual += 1
            for produto in produtos_contratados:
                for col_idx, header in enumerate(headers_produtos, 1):
                    sheet.cell(row=linha_atual, column=col_idx, value=produto.get(header, 'N/A'))
                linha_atual += 1
        excel_stream = BytesIO()
        workbook.save(excel_stream)
        excel_stream.seek(0)
        return excel_stream
    except Exception as e:
        print(f"\n[ERRO] Não foi possível salvar a planilha: {e}")
        return None